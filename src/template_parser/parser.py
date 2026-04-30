import os
import re
import tempfile
import atexit
import shutil
import logging
from collections import Counter
from statistics import median
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

from docx.enum.text import WD_ALIGN_PARAGRAPH

_EMU_PER_CM = 360000
_EMU_PER_PT = 12700

_tmp_dirs_to_cleanup = []


def _cleanup_tmp_dirs():
    for d in _tmp_dirs_to_cleanup:
        if os.path.exists(d):
            try:
                shutil.rmtree(d)
            except Exception:
                pass


atexit.register(_cleanup_tmp_dirs)


def doc_to_docx(doc_path: str) -> Optional[str]:
    word = None
    try:
        import win32com.client
        import pythoncom
        import pywintypes
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = None
        docx_path = None
        try:
            doc = word.Documents.Open(
                os.path.abspath(doc_path),
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False
            )
            try:
                tmp_dir = tempfile.mkdtemp()
                _tmp_dirs_to_cleanup.append(tmp_dir)
                docx_path = os.path.join(tmp_dir, "converted.docx")
                doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)
            finally:
                if doc:
                    try:
                        doc.Close(False)
                    except Exception:
                        pass
                    doc = None
        finally:
            if word:
                try:
                    word.ScreenUpdating = True
                except Exception:
                    pass
                word = None
        return docx_path
    except Exception as e:
        logger.warning(f"doc转docx失败: {e}")
        return None
    finally:
        import gc
        gc.collect()
        gc.collect()
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        # Clean up Word lock files if any
        try:
            lock_file = os.path.join(os.path.dirname(doc_path), "~$" + os.path.basename(doc_path))
            if os.path.exists(lock_file):
                os.remove(lock_file)
        except Exception:
            pass


def parse_template(template_path: str) -> dict:
    path = Path(template_path)
    actual_path = template_path

    if path.suffix.lower() == '.doc':
        converted = doc_to_docx(template_path)
        if converted:
            actual_path = converted
        else:
            raise ValueError(f"无法转换.doc文件: {template_path}")

    from docx import Document

    doc = Document(actual_path)

    result = {
        "page_setup": _parse_page_setup(doc),
        "header_footer": _parse_header_footer(doc),
        "content": [],
        "format_rules": {}
    }

    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p' and element in para_map:
            result["content"].append(_parse_paragraph(para_map[element]))
        elif tag == 'tbl' and element in table_map:
            result["content"].append(_parse_table(table_map[element]))

    result["format_rules"] = _extract_format_rules(result)

    return result


def _parse_page_setup(doc) -> dict:
    if doc.sections:
        section = doc.sections[0]
        return {
            "page_width_cm": round(section.page_width / _EMU_PER_CM, 2) if section.page_width else None,
            "page_height_cm": round(section.page_height / _EMU_PER_CM, 2) if section.page_height else None,
            "left_margin_cm": round(section.left_margin / _EMU_PER_CM, 2) if section.left_margin else None,
            "right_margin_cm": round(section.right_margin / _EMU_PER_CM, 2) if section.right_margin else None,
            "top_margin_cm": round(section.top_margin / _EMU_PER_CM, 2) if section.top_margin else None,
            "bottom_margin_cm": round(section.bottom_margin / _EMU_PER_CM, 2) if section.bottom_margin else None,
        }
    return {}


def _parse_header_footer(doc) -> dict:
    result = []
    if not doc.sections:
        return {"sections": result}

    for section_idx, section in enumerate(doc.sections):
        section_info: dict = {"section_index": section_idx}

        header_texts = []
        if section.header and not section.header.is_linked_to_previous:
            for para in section.header.paragraphs:
                text = para.text.strip()
                if text:
                    header_texts.append(text)

        footer_texts = []
        if section.footer and not section.footer.is_linked_to_previous:
            for para in section.footer.paragraphs:
                text = para.text.strip()
                if text:
                    footer_texts.append(text)

        first_header_texts = []
        if section.first_page_header and not section.first_page_header.is_linked_to_previous:
            for para in section.first_page_header.paragraphs:
                text = para.text.strip()
                if text:
                    first_header_texts.append(text)

        first_footer_texts = []
        if section.first_page_footer and not section.first_page_footer.is_linked_to_previous:
            for para in section.first_page_footer.paragraphs:
                text = para.text.strip()
                if text:
                    first_footer_texts.append(text)

        if header_texts:
            section_info["header_text"] = header_texts
        if footer_texts:
            section_info["footer_text"] = footer_texts
        if first_header_texts:
            section_info["first_page_header_text"] = first_header_texts
        if first_footer_texts:
            section_info["first_page_footer_text"] = first_footer_texts
        if section.different_first_page_header_footer:
            section_info["different_first_page"] = True

        if len(section_info) > 1:
            result.append(section_info)

    return {"sections": result}


def _alignment_to_str(alignment) -> Optional[str]:
    if alignment is None:
        return None
    try:
        mapping = {
            WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
            WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
            WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: "DISTRIBUTE",
        }
        return mapping.get(alignment, "LEFT")
    except:
        # 如果转换失败，返回默认值
        return "LEFT"


def _parse_paragraph(para) -> dict:
    result = {
        "type": "paragraph",
        "text": para.text,
        "style": para.style.name if para.style else None,
        "alignment": "LEFT",
        "runs": []
    }
    # 尝试获取 alignment，失败时保持默认
    try:
        result["alignment"] = _alignment_to_str(para.alignment)
    except:
        pass

    pf = para.paragraph_format
    if pf.first_line_indent:
        result["first_line_indent_cm"] = round(pf.first_line_indent / _EMU_PER_CM, 2)
    if pf.space_before:
        result["space_before_pt"] = round(pf.space_before / _EMU_PER_PT, 1)
    if pf.space_after:
        result["space_after_pt"] = round(pf.space_after / _EMU_PER_PT, 1)
    if pf.line_spacing:
        result["line_spacing"] = pf.line_spacing
    if pf.line_spacing_rule:
        result["line_spacing_rule"] = str(pf.line_spacing_rule)

    for run in para.runs:
        run_info = {"text": run.text}
        if run.font.name:
            run_info["font_name"] = run.font.name
        if run.font.size:
            run_info["font_size_pt"] = round(run.font.size / _EMU_PER_PT, 1)
        if run.font.bold:
            run_info["bold"] = True
        if run.font.italic:
            run_info["italic"] = True
        if run.font.underline:
            run_info["underline"] = True
        if run.font.color and run.font.color.rgb:
            run_info["font_color"] = str(run.font.color.rgb)
        result["runs"].append(run_info)

    return result


def _parse_table(table) -> dict:
    result = {
        "type": "table",
        "rows": len(table.rows),
        "cols": len(table.columns),
        "cells": [],
        "column_widths_cm": []
    }

    seen_tc = set()

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            seen_tc.add(tc_id)

            cell_info = {
                "row": row_idx,
                "col": col_idx,
                "text": cell.text.strip()
            }

            for span_col in range(col_idx + 1, len(row.cells)):
                try:
                    if row.cells[span_col]._tc is cell._tc:
                        cell_info["colspan"] = span_col - col_idx + 1
                        break
                except Exception:
                    pass

            for span_row in range(row_idx + 1, len(table.rows)):
                try:
                    if table.cell(span_row, col_idx)._tc is cell._tc:
                        cell_info["rowspan"] = span_row - row_idx + 1
                        break
                except Exception:
                    pass

            cell_info["paragraphs"] = []
            for p in cell.paragraphs:
                cell_info["paragraphs"].append(_parse_paragraph(p))
            result["cells"].append(cell_info)

    for col in table.columns:
        if col.width:
            result["column_widths_cm"].append(round(col.width / _EMU_PER_CM, 2))

    return result


def _extract_format_rules(parsed: dict) -> dict:
    rules = {
        "section_header": {"font_name": "", "font_size_pt": 0, "bold": False},
        "body_text": {"font_name": "", "font_size_pt": 0},
        "line_spacing_pt": 0,
        "first_line_indent_chars": 0,
        "space_before": 0,
        "space_after": 0
    }

    section_patterns = [
        r'^[一二三四五六七八九十]+、',
        r'^[（(][一二三四五六七八九十]+[）)]',
        r'^第[一二三四五六七八九十\d]+[章节章部篇]',
        r'^\d+[\.、]\s',
        r'^\d+\.\d+\s',
        r'^\d+\.\d+\.\d+\s',
        r'^Chapter\s+\d+',
        r'^Section\s+\d+',
        r'^Part\s+[IVXLCDM\d]+',
        r'^Lesson\s+\d+',
        r'^Unit\s+\d+',
        r'^Topic\s+\d+',
        r'^Module\s+\d+',
        r'^Appendix\s+[A-Z]',
        r'^(?:Introduction|Background|Purpose|Objectives?|Methods?|Results?|Discussion|Conclusions?|Summary|References?|Abstract)(?:[：:\.\s]|$)',
    ]

    font_size_groups = {}
    for elem in parsed["content"]:
        for run in elem.get("runs", []):
            fs = run.get("font_size_pt", 0)
            if fs > 0:
                fn = run.get("font_name", "")
                key = (fn, fs, run.get("bold", False))
                font_size_groups[key] = font_size_groups.get(key, 0) + 1

    body_candidate = None
    header_candidate = None
    if font_size_groups:
        sorted_by_count = sorted(font_size_groups.items(), key=lambda x: -x[1])

        for (fn, fs, bold), count in sorted_by_count:
            if not bold and fs > 0:
                body_candidate = (fn, fs, bold)
                break

        for (fn, fs, bold), count in sorted_by_count:
            if bold and fs > 0:
                header_candidate = (fn, fs, bold)
                break

        if not header_candidate and len(sorted_by_count) > 1:
            for (fn, fs, bold), count in sorted_by_count:
                if body_candidate and fs > body_candidate[1]:
                    header_candidate = (fn, fs, bold)
                    break

    if body_candidate:
        rules["body_text"]["font_name"] = body_candidate[0]
        rules["body_text"]["font_size_pt"] = body_candidate[1]
    if header_candidate:
        rules["section_header"]["font_name"] = header_candidate[0]
        rules["section_header"]["font_size_pt"] = header_candidate[1]
        rules["section_header"]["bold"] = header_candidate[2]

    section_header_styles = []
    for elem in parsed["content"]:
        if elem.get("type") != "paragraph":
            continue
        text = elem.get("text", "").strip()
        is_section = False
        for pattern in section_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                is_section = True
                break

        if not is_section:
            runs = elem.get("runs", [])
            if runs:
                is_bold = any(r.get("bold") for r in runs)
                max_fs = max((r.get("font_size_pt", 0) for r in runs), default=0)
                body_fs = rules["body_text"]["font_size_pt"]
                if is_bold and body_fs > 0 and max_fs > body_fs:
                    is_section = True

        if is_section:
            style = {"font_name": "", "font_size_pt": 0, "bold": False}
            for run in elem.get("runs", []):
                if run.get("font_name"):
                    style["font_name"] = run["font_name"]
                if run.get("font_size_pt"):
                    style["font_size_pt"] = run["font_size_pt"]
                if run.get("bold"):
                    style["bold"] = run["bold"]
            if style["font_name"] or style["font_size_pt"]:
                section_header_styles.append(style)

    if section_header_styles:
        style_keys = [(s["font_name"], s["font_size_pt"], s["bold"]) for s in section_header_styles]
        most_common_key = Counter(style_keys).most_common(1)[0][0]
        for s in section_header_styles:
            if (s["font_name"], s["font_size_pt"], s["bold"]) == most_common_key:
                rules["section_header"]["font_name"] = s["font_name"]
                rules["section_header"]["font_size_pt"] = s["font_size_pt"]
                rules["section_header"]["bold"] = s["bold"]
                break

    line_spacing_values = []
    first_line_indent_values = []
    space_before_values = []
    space_after_values = []

    for elem in parsed["content"]:
        if elem.get("type") != "paragraph":
            continue
        if elem.get("line_spacing"):
            ls = elem["line_spacing"]
            if isinstance(ls, (int, float)) and ls > 0:
                line_spacing_values.append(ls)
        if elem.get("first_line_indent_cm"):
            fi = elem["first_line_indent_cm"]
            if fi > 0:
                first_line_indent_values.append(fi)
        if elem.get("space_before_pt"):
            space_before_values.append(elem["space_before_pt"])
        if elem.get("space_after_pt"):
            space_after_values.append(elem["space_after_pt"])

    if line_spacing_values:
        rules["line_spacing_pt"] = round(median(line_spacing_values), 1)
    if first_line_indent_values:
        body_fs = rules["body_text"]["font_size_pt"] or 12
        avg_indent_cm = median(first_line_indent_values)
        char_width_cm = body_fs * 0.035
        rules["first_line_indent_chars"] = max(0, round(avg_indent_cm / char_width_cm)) if char_width_cm > 0 else 0
    if space_before_values:
        rules["space_before"] = round(median(space_before_values), 1)
    if space_after_values:
        rules["space_after"] = round(median(space_after_values), 1)

    if not rules["body_text"]["font_name"]:
        rules["body_text"]["font_name"] = ""
    if not rules["body_text"]["font_size_pt"]:
        rules["body_text"]["font_size_pt"] = 0
    if not rules["section_header"]["font_name"]:
        rules["section_header"]["font_name"] = ""
    if not rules["section_header"]["font_size_pt"]:
        rules["section_header"]["font_size_pt"] = 0
    if not rules["section_header"]["bold"]:
        rules["section_header"]["bold"] = False

    return rules
