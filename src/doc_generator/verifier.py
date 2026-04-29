import os
import json
import tempfile
import logging
from typing import Optional

from docx import Document

from src.template_parser.parser import doc_to_docx

logger = logging.getLogger(__name__)


def verify_format(template_path: str, generated_path: str, output_path: Optional[str] = None,
                   profile: Optional[dict] = None) -> dict:
    template_docx = _ensure_docx(template_path)
    generated_docx = _ensure_docx(generated_path)

    try:
        template_doc = Document(template_docx)
        generated_doc = Document(generated_docx)

        results = {
            "page_setup": _verify_page_setup(template_doc, generated_doc),
            "paragraphs": _verify_paragraphs(template_doc, generated_doc, profile),
            "tables": _verify_tables(template_doc, generated_doc),
            "passed": True,
            "issues": [],
            "requirement_warnings": []
        }

        for category, checks in results.items():
            if category in ["passed", "issues", "requirement_warnings"]:
                continue
            if isinstance(checks, dict):
                for key, value in checks.items():
                    if value is False:
                        results["passed"] = False
                        results["issues"].append(f"{category}.{key}: 不一致")
            elif isinstance(checks, list):
                for i, check in enumerate(checks):
                    if isinstance(check, dict):
                        for key, value in check.items():
                            if value is False and key not in ("text_match", "match_type"):
                                results["passed"] = False
                                results["issues"].append(f"{category}[{i}].{key}: 不一致")

        if profile:
            req_warnings = _check_requirements(generated_doc, profile)
            results["requirement_warnings"] = req_warnings

        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=2)

        return results
    finally:
        if template_docx != template_path and os.path.exists(template_docx):
            try:
                os.remove(template_docx)
            except Exception:
                pass
        if generated_docx != generated_path and os.path.exists(generated_docx):
            try:
                os.remove(generated_docx)
            except Exception:
                pass


def _check_requirements(generated_doc, profile: dict) -> list:
    warnings = []
    sections = profile.get("sections", [])

    section_titles = []
    for sec in sections:
        section_titles.append(sec.get("title", ""))

    para_section_map = _map_paras_to_sections(generated_doc, section_titles)

    for sec in sections:
        title = sec.get("title", "")
        requirements = sec.get("requirements", [])
        if not requirements:
            continue

        section_para_indices = para_section_map.get(title, [])

        for req in requirements:
            req_type = req.get("type", "")
            desc = req.get("description", "")
            value = req.get("value", "")

            if req_type == "min_count" and value:
                try:
                    min_count = int(value)
                    content_paras = [i for i in section_para_indices if i not in _get_title_indices(generated_doc, section_titles)]
                    actual_count = 0
                    for idx in content_paras:
                        if idx < len(generated_doc.paragraphs):
                            text = generated_doc.paragraphs[idx].text.strip()
                            if text:
                                actual_count += 1
                    if actual_count < min_count:
                        warnings.append(f"章节「{title}」: 约束「{desc}」可能未满足 — 检测到{actual_count}个非空段落，要求不少于{min_count}个")
                except (ValueError, IndexError):
                    pass

            elif req_type == "font" and value:
                font_found = False
                for idx in section_para_indices:
                    if idx < len(generated_doc.paragraphs):
                        for run in generated_doc.paragraphs[idx].runs:
                            if run.font.name and value.lower() in run.font.name.lower():
                                font_found = True
                                break
                            try:
                                rPr = run._element.rPr
                                if rPr is not None:
                                    east_asia = rPr.rFonts.get(
                                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                                    if east_asia and value.lower() in east_asia.lower():
                                        font_found = True
                                        break
                            except Exception:
                                pass
                        if font_found:
                            break
                if not font_found:
                    warnings.append(f"章节「{title}」: 约束「{desc}」可能未满足 — 未检测到字体「{value}」")

            elif req_type == "table_structure":
                has_table = False
                for idx in section_para_indices:
                    if idx < len(generated_doc.paragraphs):
                        p_elem = generated_doc.paragraphs[idx]._element
                        next_sib = p_elem.getnext()
                        if next_sib is not None and next_sib.tag.endswith('}tbl'):
                            has_table = True
                            break
                if not has_table:
                    warnings.append(f"章节「{title}」: 约束「{desc}」可能未满足 — 未检测到表格")

            elif req_type in ("content", "format", "other"):
                warnings.append(f"章节「{title}」: 提醒约束「{desc}」— 请人工确认是否满足")

    return warnings


def _map_paras_to_sections(doc, section_titles: list) -> dict:
    result = {t: [] for t in section_titles}
    current_section = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text in section_titles:
            current_section = text
            result[current_section].append(i)
            continue
        if current_section:
            result[current_section].append(i)

    return result


def _get_title_indices(doc, section_titles: list) -> list:
    indices = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() in section_titles:
            indices.append(i)
    return indices


def _ensure_docx(file_path: str) -> str:
    if file_path.lower().endswith('.docx'):
        return file_path

    converted = doc_to_docx(file_path)
    if converted:
        return converted

    raise ValueError(f"无法转换.doc文件为.docx格式: {file_path}。请确保系统已安装Microsoft Word。")


def _verify_page_setup(template_doc, generated_doc) -> dict:
    t_sec = template_doc.sections[0] if template_doc.sections else None
    g_sec = generated_doc.sections[0] if generated_doc.sections else None

    if not t_sec or not g_sec:
        return {"has_sections": t_sec is not None and g_sec is not None}

    return {
        "page_width": _compare_emu(t_sec.page_width, g_sec.page_width),
        "page_height": _compare_emu(t_sec.page_height, g_sec.page_height),
        "left_margin": _compare_emu(t_sec.left_margin, g_sec.left_margin),
        "right_margin": _compare_emu(t_sec.right_margin, g_sec.right_margin),
        "top_margin": _compare_emu(t_sec.top_margin, g_sec.top_margin),
        "bottom_margin": _compare_emu(t_sec.bottom_margin, g_sec.bottom_margin),
    }


def _compare_emu(a, b, tolerance: int = 10000) -> bool:
    if a is None or b is None:
        return a is None and b is None
    return abs(a - b) < tolerance


def _verify_paragraphs(template_doc, generated_doc, profile: Optional[dict] = None) -> list:
    results = []

    if profile:
        _verify_landmarks(template_doc, generated_doc, profile, results)
        _verify_body_format(generated_doc, profile, results)
    else:
        _verify_by_text_match(template_doc, generated_doc, results)

    return results


def _find_para_by_text(doc, text: str):
    for para in doc.paragraphs:
        if para.text.strip() == text.strip():
            return para
    return None


def _compare_para_format(t_para, g_para, label: str) -> dict:
    result = {
        "label": label,
        "text_match": t_para.text.strip() == g_para.text.strip(),
    }
    result["alignment_match"] = t_para.alignment == g_para.alignment

    t_runs = t_para.runs
    g_runs = g_para.runs
    if t_runs and g_runs:
        result["font_name_match"] = _compare_font_name(t_runs[0], g_runs[0])
        result["font_size_match"] = _compare_font_size(t_runs[0], g_runs[0])
        result["bold_match"] = t_runs[0].font.bold == g_runs[0].font.bold
        result["italic_match"] = t_runs[0].font.italic == g_runs[0].font.italic
        result["underline_match"] = t_runs[0].font.underline == g_runs[0].font.underline

    return result


def _get_profile_landmarks(profile: dict) -> dict:
    titles = set()
    for sec in profile.get("sections", []):
        t = sec.get("title", "").strip()
        if t:
            titles.add(t)
    labels = set()
    for field in profile.get("cover_page", {}).get("fields", []):
        lbl = field.get("label", "").strip()
        if lbl:
            labels.add(lbl)
    return {"section_titles": titles, "cover_labels": labels}


def _verify_landmarks(template_doc, generated_doc, profile: dict, results: list):
    landmarks = _get_profile_landmarks(profile)

    for title in landmarks["section_titles"]:
        t_para = _find_para_by_text(template_doc, title)
        g_para = _find_para_by_text(generated_doc, title)
        if t_para and g_para:
            r = _compare_para_format(t_para, g_para, f"section_title: {title[:30]}")
            r["match_type"] = "landmark_section"
            results.append(r)
        elif g_para:
            results.append({
                "match_type": "landmark_section",
                "label": f"section_title: {title[:30]}",
                "text_match": True,
                "note": "section title found in generated doc only (template match skipped)"
            })
        elif t_para:
            results.append({
                "match_type": "landmark_section",
                "label": f"section_title: {title[:30]}",
                "text_match": False,
                "note": "section title found in template only — may have been removed"
            })

    for label in landmarks["cover_labels"]:
        t_para = _find_para_by_text(template_doc, label)
        g_para = _find_para_by_text(generated_doc, label)
        if t_para and g_para:
            r = _compare_para_format(t_para, g_para, f"cover_label: {label[:30]}")
            r["match_type"] = "landmark_cover"
            results.append(r)
        elif g_para:
            results.append({
                "match_type": "landmark_cover",
                "label": f"cover_label: {label[:30]}",
                "text_match": True,
                "note": "cover label found in generated doc only (template match skipped)"
            })


def _verify_body_format(generated_doc, profile: dict, results: list):
    rules = profile.get("format_rules", {})
    body = rules.get("body_text", {})
    expected_font = body.get("font_name", "")
    expected_size = body.get("font_size_pt", 0)

    if not expected_font and not expected_size:
        return

    section_titles = set()
    for sec in profile.get("sections", []):
        t = sec.get("title", "").strip()
        if t:
            section_titles.add(t)

    sampled = 0
    for para in generated_doc.paragraphs:
        text = para.text.strip()
        if not text or text in section_titles:
            continue
        if not para.runs:
            continue
        if sampled >= 5:
            break

        run = para.runs[0]
        mismatches = {}
        if expected_font:
            fn = run.font.name
            if fn and fn != expected_font:
                try:
                    rPr = run._element.rPr
                    if rPr is not None:
                        ea = rPr.rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                        if ea and ea != expected_font:
                            mismatches["body_font_name"] = f"expected={expected_font}, got=(name={fn}, eastAsia={ea})"
                except Exception:
                    if fn != expected_font:
                        mismatches["body_font_name"] = f"expected={expected_font}, got={fn}"
            elif fn and fn == expected_font:
                pass  # match

        if expected_size and run.font.size:
            actual_pt = run.font.size / 12700
            if abs(actual_pt - expected_size) > 1.0:
                mismatches["body_font_size"] = f"expected={expected_size}pt, got={actual_pt:.1f}pt"

        if mismatches:
            mismatches["match_type"] = "body_sample"
            mismatches["label"] = f"body_sample: {text[:30]}"
            mismatches["text_match"] = True
            results.append(mismatches)

        sampled += 1


def _verify_by_text_match(template_doc, generated_doc, results: list):
    t_paras = template_doc.paragraphs
    g_paras = generated_doc.paragraphs

    g_para_map = {}
    for i, g_para in enumerate(g_paras):
        text = g_para.text.strip()
        if text:
            g_para_map.setdefault(text, []).append(i)

    matched_g_indices = set()

    for i, t_para in enumerate(t_paras):
        t_text = t_para.text.strip()
        if not t_text:
            continue

        g_indices = g_para_map.get(t_text, [])
        g_idx = None
        for gi in g_indices:
            if gi not in matched_g_indices:
                g_idx = gi
                matched_g_indices.add(gi)
                break

        if g_idx is None:
            results.append({
                "template_index": i,
                "template_text": t_text[:50],
                "generated_text": "(not found)",
                "text_match": False,
            })
            continue

        g_para = g_paras[g_idx]
        result = {
            "template_index": i,
            "generated_index": g_idx,
            "template_text": t_text[:50],
            "generated_text": g_para.text.strip()[:50],
            "text_match": True,
        }

        result["alignment_match"] = t_para.alignment == g_para.alignment

        t_runs = t_para.runs
        g_runs = g_para.runs

        if t_runs and g_runs:
            result["font_name_match"] = _compare_font_name(t_runs[0], g_runs[0])
            result["font_size_match"] = _compare_font_size(t_runs[0], g_runs[0])
            result["bold_match"] = t_runs[0].font.bold == g_runs[0].font.bold
            result["italic_match"] = t_runs[0].font.italic == g_runs[0].font.italic
            result["underline_match"] = t_runs[0].font.underline == g_runs[0].font.underline

        results.append(result)


def _compare_font_name(run_a, run_b) -> bool:
    a_name = run_a.font.name
    b_name = run_b.font.name
    if a_name == b_name:
        return True
    try:
        a_east = run_a._element.rPr.rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
        b_east = run_b._element.rPr.rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
        return a_east == b_east
    except Exception:
        return False


def _compare_font_size(run_a, run_b, tolerance: int = 12700) -> bool:
    a_size = run_a.font.size
    b_size = run_b.font.size
    if a_size is None or b_size is None:
        return a_size is None and b_size is None
    return abs(a_size - b_size) < tolerance


def _verify_tables(template_doc, generated_doc) -> list:
    results = []
    t_tables = template_doc.tables
    g_tables = generated_doc.tables

    min_len = min(len(t_tables), len(g_tables))
    for i in range(min_len):
        t_table = t_tables[i]
        g_table = g_tables[i]

        result = {
            "index": i,
            "rows_match": len(t_table.rows) == len(g_table.rows),
            "cols_match": len(t_table.columns) == len(g_table.columns),
        }

        results.append(result)

    return results
