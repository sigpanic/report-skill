import os
import re
import shutil
import tempfile
import logging
from pathlib import Path
from typing import Optional
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.text.paragraph import Paragraph

from src.template_parser.parser import doc_to_docx

logger = logging.getLogger(__name__)

_HINT_COLOR_RANGES = [
    (200, 0, 0), (255, 0, 0), (255, 51, 51),
    (255, 80, 80), (200, 50, 50),
    (128, 128, 128), (160, 160, 160),
    (100, 100, 200), (0, 0, 200),
]


def generate_report(
    template_path: str,
    output_path: str,
    profile: dict,
    field_values: dict,
    sections: list,
    result_images: Optional[list] = None
) -> str:
    path = Path(template_path)
    work_docx = None
    template_is_doc = path.suffix.lower() == '.doc'
    tmp_dir = None
    tmp_dir2 = None

    if template_is_doc:
        converted = doc_to_docx(template_path)
        if converted:
            work_docx = converted
        else:
            raise ValueError(f"无法转换.doc文件: {template_path}")
    else:
        tmp_dir = tempfile.mkdtemp()
        work_docx = os.path.join(tmp_dir, "template_copy.docx")
        shutil.copy2(template_path, work_docx)

    try:
        doc = Document(work_docx)

        _fill_cover_fields(doc, profile, field_values)
        _fill_table_fields(doc, profile, field_values)
        _fill_sections(doc, profile, sections, result_images)
        _remove_annotations(doc, profile)
        _ensure_header_footer(doc, profile)

        out_path = Path(output_path)
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

        if template_is_doc:
            output_path = str(out_path.with_suffix('.doc'))
            tmp_dir2 = tempfile.mkdtemp()
            docx_output = os.path.join(tmp_dir2, "temp_output.docx")
            doc.save(docx_output)
            doc_to_save = _docx_to_doc(docx_output, output_path)
            if doc_to_save:
                return doc_to_save
            docx_output_path = str(out_path.with_suffix('.docx'))
            shutil.copy2(docx_output, docx_output_path)
            logger.warning("docx转doc失败，已保存为docx格式: %s", docx_output_path)
            return docx_output_path

        doc.save(output_path)
        return output_path
    finally:
        for d in [tmp_dir, tmp_dir2]:
            if d and os.path.exists(d):
                try:
                    shutil.rmtree(d)
                except Exception:
                    pass


def _docx_to_doc(docx_path: str, doc_output_path: str) -> Optional[str]:
    doc = None
    word = None
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        result = None
        try:
            doc = word.Documents.Open(os.path.abspath(docx_path))
            try:
                doc.SaveAs2(os.path.abspath(doc_output_path), FileFormat=0)
                result = doc_output_path
            finally:
                if doc:
                    try:
                        doc.Close(False)
                    except Exception:
                        pass
                    doc = None
        finally:
            if word:
                word = None
        return result
    except Exception as e:
        logger.warning("docx转doc失败: %s", e)
        return None
    finally:
        try:
            import gc
            gc.collect()
            gc.collect()
        except Exception:
            pass
        try:
            import pythoncom
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _fill_cover_fields(doc, profile: dict, field_values: dict):
    cover = profile.get("cover_page", {})

    all_paras = list(doc.paragraphs)

    for para in all_paras:
        text = para.text.strip()

        for field in cover.get("fields", []):
            label = field.get("label", "")
            key = field.get("key", "")
            value = field_values.get(key, field.get("default", ""))

            if not value:
                continue

            if _fuzzy_match_label(text, label):
                _replace_cover_field_run_level(para, label, value, field)
                break

            if not label and field.get("default") and text == field["default"]:
                for run in para.runs:
                    if run.text.strip():
                        run.text = value
                break


def _fuzzy_match_label(text: str, label: str) -> bool:
    if not label:
        return False
    text_clean = re.sub(r'\s+', '', text).replace('：', ':')
    label_clean = re.sub(r'\s+', '', label).replace('：', ':')
    return text_clean.startswith(label_clean)


def _replace_cover_field_run_level(para, label: str, value: str, field: dict):
    if not para.runs:
        return

    label_runs = []
    value_runs = []
    found_label_end = False

    for run in para.runs:
        if not found_label_end:
            label_runs.append(run)
            if label in run.text or run.text.rstrip().endswith(("：", ":")):
                found_label_end = True
        else:
            value_runs.append(run)

    if not value_runs:
        label_text = "".join(r.text for r in label_runs)
        clean_label = label_text.rstrip("：: ").rstrip()
        found_label_end = False
        for run in label_runs:
            run_text_clean = run.text.rstrip("：: ").rstrip()
            if not found_label_end:
                if run_text_clean == clean_label or clean_label.endswith(run_text_clean):
                    found_label_end = True
                    continue
            if found_label_end:
                value_runs.append(run)

    if not value_runs:
        _replace_para_simple(para, label, value, field)
        return

    total_value_len = sum(len(r.text) for r in value_runs)

    new_value = value
    padding_needed = max(0, total_value_len - len(new_value))
    padded_value = new_value + " " * padding_needed

    char_idx = 0
    for i, run in enumerate(value_runs):
        run_len = len(run.text)
        if i < len(value_runs) - 1:
            run.text = padded_value[char_idx:char_idx + run_len]
            char_idx += run_len
        else:
            run.text = padded_value[char_idx:]


def _replace_para_simple(para, label: str, value: str, field: dict):
    has_underline = field.get("type") == "text_with_underline" or any(
        r.font.underline for r in para.runs
    )

    for run in para.runs:
        run.text = ""

    para.runs[0].text = label + value
    if has_underline:
        para.runs[0].font.underline = True


def _fill_table_fields(doc, profile: dict, field_values: dict):
    tables_profile = profile.get("tables", [])

    profile_matched = set()
    for table_idx, table in enumerate(doc.tables):
        best_match = -1
        best_score = -1
        for pi, tp in enumerate(tables_profile):
            if pi in profile_matched:
                continue
            if len(table.rows) == tp.get("rows", 0) and len(table.columns) == tp.get("cols", 0):
                score = 1
                for field in tp.get("fields", []):
                    cell_pos = field.get("cell", "").split(",")
                    if len(cell_pos) == 2:
                        try:
                            r, c = int(cell_pos[0]), int(cell_pos[1])
                            cell_text = table.cell(r, c).text.strip()
                            if field.get("label", "") and cell_text == field["label"]:
                                score += 2
                            elif cell_text:
                                score += 1
                        except (IndexError, ValueError):
                            pass
                if score > best_score:
                    best_score = score
                    best_match = pi

        if best_match < 0:
            if table_idx < len(tables_profile):
                best_match = table_idx
            else:
                continue

        profile_matched.add(best_match)
        table_profile = tables_profile[best_match]

        for field in table_profile.get("fields", []):
            key = field.get("key", "")
            value = field_values.get(key, field.get("default", ""))

            if not value:
                continue

            cell_pos = field.get("cell", "").split(",")
            if len(cell_pos) != 2:
                continue

            try:
                row, col = int(cell_pos[0]), int(cell_pos[1])
                cell = table.cell(row, col)
                _fill_cell_preserving_style(cell, value, field)
            except (IndexError, ValueError) as e:
                logger.warning("表格字段填充失败 key=%s cell=%s: %s", key, field.get("cell", ""), e)
            except Exception as e:
                logger.warning("表格字段填充异常 key=%s: %s", key, e)


def _fill_cell_preserving_style(cell, value: str, field: dict):
    if not cell.paragraphs:
        return

    p = cell.paragraphs[0]
    original_alignment = p.alignment

    if p.runs:
        hint_runs = [r for r in p.runs if r.font.italic or _is_hint_run(r)]
        if hint_runs:
            for r in hint_runs:
                r.text = ""
            hint_runs[0].text = value
            hint_runs[0].font.italic = False
            try:
                hint_runs[0].font.color.rgb = None
            except Exception:
                pass
            if original_alignment is not None:
                p.alignment = original_alignment
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return

        for r in p.runs:
            r.text = ""
        p.runs[0].text = value
    else:
        run = p.add_run(value)
        style = field.get("style", {})
        font_name = style.get("font_name", "")
        font_size = style.get("font_size_pt", 0)
        if font_name:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            run.font.size = Pt(font_size)

    if original_alignment is not None:
        p.alignment = original_alignment
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _is_hint_run(run) -> bool:
    try:
        if run.font.color and run.font.color.rgb:
            rgb_str = str(run.font.color.rgb).upper()
            if len(rgb_str) == 6:
                r_val = int(rgb_str[0:2], 16)
                g_val = int(rgb_str[2:4], 16)
                b_val = int(rgb_str[4:6], 16)
                for hr, hg, hb in _HINT_COLOR_RANGES:
                    if abs(r_val - hr) <= 40 and abs(g_val - hg) <= 40 and abs(b_val - hb) <= 40:
                        return True
    except Exception:
        pass
    return False


def _fill_sections(doc, profile: dict, sections: list, result_images: Optional[list] = None):
    format_rules = profile.get("format_rules", {})
    section_map = {}
    for sec in sections:
        section_map[sec["title"]] = sec

    for para in doc.paragraphs:
        text = para.text.strip()
        if text in section_map:
            sec = section_map[text]
            sec_profile = _find_section_profile(profile, text)
            content_style = sec_profile.get("content_style") if sec_profile else None
            requirements = sec_profile.get("requirements", []) if sec_profile else []
            _insert_section_content(
                para, sec.get("content", ""), sec.get("images", []),
                sec.get("tables", []), format_rules, content_style, requirements
            )

    if result_images:
        last_section_para = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if text in section_map:
                last_section_para = para
        if last_section_para:
            _insert_result_images_after(last_section_para, result_images, format_rules)


def _find_section_profile(profile: dict, title: str) -> dict:
    for sec in profile.get("sections", []):
        if sec.get("title") == title:
            return sec
    return {}


def _insert_result_images_after(title_para, image_paths: list, format_rules: dict):
    body_style = format_rules.get("body_text", {})
    font_name = body_style.get("font_name", "")

    insert_after = title_para._element
    for img_path in image_paths:
        if not os.path.exists(img_path):
            continue
        img_para_xml = parse_xml(
            f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>'
        )
        img_para_xml.set(qn('w:customXml'), "ins")
        insert_after.addnext(img_para_xml)
        insert_after = img_para_xml

        img_para = Paragraph(img_para_xml, title_para._element.getparent())
        img_run = img_para.add_run()
        img_width = format_rules.get("image_width_cm") or 12.0
        if img_width <= 0:
            img_width = 12.0
        img_run.add_picture(img_path, width=Cm(img_width))


def _insert_section_content(title_para, content: str, images: list, content_tables: list,
                            format_rules: dict, content_style: Optional[dict] = None, requirements: Optional[list] = None):
    body_style = format_rules.get("body_text", {})
    font_name = content_style.get("font_name", "") if content_style and content_style.get("font_name") else body_style.get("font_name", "")
    font_size_pt = content_style.get("font_size_pt", 0) if content_style and content_style.get("font_size_pt") else body_style.get("font_size_pt", 0)
    line_spacing_pt = format_rules.get("line_spacing_pt", 0)
    indent_chars = format_rules.get("first_line_indent_chars", 0)
    indent_cm = indent_chars * 0.42 if indent_chars else 0
    space_before_pt = format_rules.get("space_before", 0)
    space_after_pt = format_rules.get("space_after", 0)
    font_name_safe = xml_escape(font_name, {'"': '&quot;'})
    is_italic = content_style.get("italic", False) if content_style else False
    is_underline = content_style.get("underline", False) if content_style else False
    alignment = (content_style.get("alignment", "") or "").upper() if content_style else ""

    insert_after = title_para._element
    next_sib = insert_after.getnext()
    while next_sib is not None:
        next_tag = next_sib.tag.split('}')[-1] if '}' in next_sib.tag else next_sib.tag
        if next_tag != 'p':
            break
        next_para = Paragraph(next_sib, title_para._element.getparent())
        is_annotation = False
        for run in next_para.runs:
            if run.font.italic or _is_hint_run(run):
                is_annotation = True
                break
        if is_annotation:
            insert_after = next_sib
            next_sib = next_sib.getnext()
        else:
            break

    if content:
        content_lines = content.split('\n')
        for line in content_lines:
            sz_val = int(font_size_pt * 2) if font_size_pt else 24
            indent_xml = ""
            if indent_cm:
                indent_twips = round(indent_cm / 2.54 * 1440)
                indent_xml = f'<w:ind w:firstLine="{indent_twips}"/>'
            new_para_xml = parse_xml(
                f'<w:p {nsdecls("w")}>'
                f'<w:pPr>'
                f'{indent_xml}'
                f'<w:rPr><w:rFonts w:eastAsia="{font_name_safe}"/><w:sz w:val="{sz_val}"/></w:rPr>'
                f'</w:pPr>'
                f'</w:p>'
            )
            new_para_xml.set(qn('w:customXml'), "ins")
            insert_after.addnext(new_para_xml)
            insert_after = new_para_xml

            new_p = Paragraph(new_para_xml, title_para._element.getparent())

            if line.strip():
                run = new_p.add_run(line)
                if font_name:
                    run.font.name = font_name
                    rPr = run._element.rPr
                    if rPr is not None:
                        rPr.rFonts.set(qn('w:eastAsia'), font_name)
                if font_size_pt:
                    run.font.size = Pt(font_size_pt)
                if is_italic:
                    run.font.italic = True
                if is_underline:
                    run.font.underline = True

                pf = new_p.paragraph_format
                if indent_cm:
                    pf.first_line_indent = Cm(indent_cm)
                if line_spacing_pt:
                    pf.line_spacing = Pt(line_spacing_pt)
                pf.space_before = Pt(space_before_pt)
                pf.space_after = Pt(space_after_pt)
                if alignment == "CENTER":
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == "RIGHT":
                    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == "LEFT":
                    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif alignment == "JUSTIFY":
                    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                elif alignment == "DISTRIBUTE":
                    pf.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

    if content_tables:
        for table_data in content_tables:
            col_widths = table_data.get("column_widths_cm") or table_data.get("col_widths")
            tbl_element = _create_table_element(table_data, font_name, font_size_pt,
                                                  format_rules.get("table_header_bg_color", ""),
                                                  column_widths_cm=col_widths)
            insert_after.addnext(tbl_element)
            insert_after = tbl_element

    if images:
        for img_path in images:
            if os.path.exists(img_path):
                img_para_xml = parse_xml(
                    f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>'
                )
                img_para_xml.set(qn('w:customXml'), "ins")
                insert_after.addnext(img_para_xml)
                insert_after = img_para_xml

                img_para = Paragraph(img_para_xml, title_para._element.getparent())
                img_run = img_para.add_run()
                img_width = format_rules.get("image_width_cm") or 12.0
                if img_width <= 0:
                    img_width = 12.0
                img_run.add_picture(img_path, width=Cm(img_width))


def _remove_annotations(doc, profile: dict):
    annotation_patterns = profile.get("annotation_patterns", [])
    removal_patterns = profile.get("removal_patterns", [])

    paras_to_remove = []
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    for para in all_paras:
        p_element = para._element
        if p_element.get(qn('w:customXml')) == "ins":
            continue

        text = para.text.strip()

        should_remove = False
        for pattern in annotation_patterns:
            if pattern in text:
                should_remove = True
                break

        if not should_remove:
            for pattern in removal_patterns:
                try:
                    if re.search(pattern, text):
                        should_remove = True
                        break
                except re.error:
                    if pattern in text:
                        should_remove = True
                        break

        if should_remove:
            paras_to_remove.append(para)

    for para in paras_to_remove:
        p_element = para._element
        if p_element is not None and p_element.getparent() is not None:
            p_element.getparent().remove(p_element)


def _create_table_element(table_data: dict, font_name: str, font_size_pt: float, header_bg_color: str = "",
                          column_widths_cm: Optional[list] = None):
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)

    if num_cols == 0:
        return parse_xml(f'<w:p {nsdecls("w")}></w:p>')

    ns = nsdecls("w")
    sz = int(font_size_pt * 2) if font_size_pt else 24
    fn = xml_escape(font_name, {'"': '&quot;'})

    tbl_xml = f'<w:tbl {ns}>'
    tbl_xml += f'<w:tblPr><w:tblBorders>'
    tbl_xml += f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    tbl_xml += f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    tbl_xml += f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    tbl_xml += f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    tbl_xml += f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    tbl_xml += f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    tbl_xml += f'</w:tblBorders><w:jc w:val="center"/></w:tblPr>'

    if column_widths_cm:
        tbl_xml += '<w:tblGrid>'
        for w_cm in column_widths_cm:
            w_twips = int(w_cm / 2.54 * 1440)
            tbl_xml += f'<w:gridCol w:w="{w_twips}"/>'
        tbl_xml += '</w:tblGrid>'

    if headers:
        tbl_xml += '<w:tr>'
        for h in headers:
            h_escaped = xml_escape(str(h))
            bg = f'<w:tc><w:tcPr><w:shd w:val="clear" w:color="auto" w:fill="{header_bg_color}"/></w:tcPr>' if header_bg_color else '<w:tc><w:tcPr/>'
            tbl_xml += bg
            tbl_xml += f'<w:p><w:pPr><w:jc w:val="center"/><w:rPr><w:rFonts w:eastAsia="{fn}"/><w:sz w:val="{sz}"/><w:b/></w:rPr></w:pPr>'
            tbl_xml += f'<w:r><w:rPr><w:rFonts w:eastAsia="{fn}"/><w:sz w:val="{sz}"/><w:b/></w:rPr><w:t>{h_escaped}</w:t></w:r>'
            tbl_xml += f'</w:p></w:tc>'
        tbl_xml += '</w:tr>'

    for row in rows:
        tbl_xml += '<w:tr>'
        for cell in row:
            cell_escaped = xml_escape(str(cell))
            tbl_xml += f'<w:tc><w:p><w:pPr><w:jc w:val="center"/><w:rPr><w:rFonts w:eastAsia="{fn}"/><w:sz w:val="{sz}"/></w:rPr></w:pPr>'
            tbl_xml += f'<w:r><w:rPr><w:rFonts w:eastAsia="{fn}"/><w:sz w:val="{sz}"/></w:rPr><w:t>{cell_escaped}</w:t></w:r>'
            tbl_xml += f'</w:p></w:tc>'
        tbl_xml += '</w:tr>'

    tbl_xml += '</w:tbl>'
    return parse_xml(tbl_xml)


def _ensure_header_footer(doc, profile: dict):
    hf = profile.get("header_footer", {})
    if not hf:
        return

    if not doc.sections:
        return

    for section_info in hf.get("sections", []):
        idx = section_info.get("section_index", 0)
        if idx >= len(doc.sections):
            continue
        section = doc.sections[idx]

        if section_info.get("different_first_page"):
            section.different_first_page_header_footer = True
